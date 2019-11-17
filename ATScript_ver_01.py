import PortSerial
import os
import time
from array import array
import matplotlib.pyplot as plt
import statistics
import numpy as np
from docx import Document
from docx.shared import Inches
from tkinter import *
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from email.utils import formatdate
# import logging
# import logging.handlers


class Main:

    def __init__(self, name_tune, params, num_iterations, comm, global_path,
                 add_params, mech_setup, advance_traj, advance_options, mail_param, smart_factor):

        self.nameComport = comm

        PortSerial.port("password he110", self.nameComport, 1)

        if int(PortSerial.port("opmode", self.nameComport, 0)) != 8:
            PortSerial.port("opmode 8", self.nameComport, 1)

        PortSerial.port("en", self.nameComport, 1)
        if int(PortSerial.port("active", self.nameComport, 0)) != 1:
            print("Drive not Enable")

        self.err_code = ['Unexpected Fault', 'Unexpected Hold Fault', 'Bundle AT Not supported', 'DECSTOP too low',
                         'Motor Was Not Recognized', 'ICMD Sat Fault', 'Foldback Occured', 'Unexpected drive warning',
                         'LMJR Estimation Failed', 'Low Acc or Velocity', 'Cycle identify timeout',
                         'Cycle count value is high', 'Trajectory Too Long', 'Stopped']

        self.name_tune = name_tune
        self.numOfIterations = num_iterations
        self.parameters = params
        self.global_path = global_path
        self.mech_setup = mech_setup
        self.sedrive = False
        self.smart_factor = smart_factor
        self.time_rec = 2000
        self.sample_value = 1

        self.mail_cb = mail_param[0]
        self.address = mail_param[1]
        self.mail_pass = mail_param[2]

        self.factor_ptpvcmd_pe = 0.01
        self.factor_ptpvcmd_icmd = 0.01
        self.factorPE = 1

        self.array_results = array('i', [])  # Array iterations
        self.traj_array = ["0", "0", "0", "0"]
        self.num_of_rec_objects = 3

        self.acc = advance_traj[0]
        self.distance = advance_traj[1]
        self.vel = advance_traj[2]

        self.avmode = "0"
        self.ffmode = "0"
        self.igrav = "0"

        self.max_velocity = 400

        if advance_options[0] is True:
            print("AV mode selected")
            self.avmode = "6"
        if advance_options[1] is True:
            print("Acc Feed Forward mode selected")
            self.ffmode = "2"
        if advance_options[2] is True:
            print("Igrav selected")
            self.igrav = "1"

        self.feedbacktype_val = int(PortSerial.port("feedbacktype", self.nameComport, 0))
        if (self.feedbacktype_val is 12) or (self.feedbacktype_val is 19):  # If feedback SensAr set factor for PE=0.002
            self.factorPE = 0.002

        self.directory = self.glob_dir_make(self.global_path)
        self.dir_plot = self.plot_dir_make()
        self.dir_iter = ""

        self.MatrixParams = [[0 for e in range(int(len(self.parameters)))]
                             for v in range(self.numOfIterations)]

        w, h = num_iterations * 2, self.time_rec
        self.MatrixPEAll = [[0 for k in range(w)] for d in range(h)]
        self.MatrixPEAll_deg = [[0 for k in range(w)] for d in range(h)]
        self.MatrixICMDAll = [[0 for e in range(w)] for v in range(h)]
        self.MatrixPEDec = [[0 for k in range(w)] for d in range(h)]
        self.MatrixPEDec_deg = [[0 for k in range(w)] for d in range(h)]
        self.MatrixICMDDec = [[0 for e in range(w)] for v in range(h)]
        self.Matrix_add_name_params = add_params
        self.Matrix_add_out_params = [[0 for e in range(len(self.Matrix_add_name_params))]
                                      for v in range(self.numOfIterations)]

    @staticmethod
    def glob_dir_make(global_path):

        directory = global_path + '\\DataCollect' + str(int(time.time()))
        if not os.path.exists(directory):
            os.makedirs(directory)
        return directory

    def plot_dir_make(self):
        dir_plot = self.directory + '\\Plots'
        if not os.path.exists(dir_plot):
            os.makedirs(dir_plot)
        return dir_plot

    def calc_factor(self, MatrixPE, MatrixICMD):

        max_val_pe = 1
        max_val_icmd = 1
        min_val_pe = 0
        min_val_icmd = 0
        max_val_ptpvcmd_vs_pe = 1

        for ind in range(0, self.numOfIterations * 2, 2):

            calc_max_val_pe = max(np.multiply(self.factorPE, MatrixPE[ind + 1][:]))  # calculation max value of
            # PE for each iteration
            calc_min_val_pe = min(np.multiply(self.factorPE, MatrixPE[ind + 1][:]))  # calculation min value of
            # PE for each iteration
            calc_max_val_ptpvcmd = max(MatrixPE[ind][:])  # calculation max value of PTPVCMD for each iteration

            calc_max_val_icmd = max(MatrixICMD[ind + 1][:])  # calculation max value of ICMD for each iteration
            calc_min_val_icmd = min(MatrixPE[ind + 1][:])  # calculation min value of ICMD for each iteration

            if calc_max_val_pe > max_val_pe:
                max_val_pe = calc_max_val_pe

            if (calc_min_val_pe < min_val_pe) and (calc_min_val_pe > 0):
                min_val_pe = calc_min_val_pe

            if calc_max_val_ptpvcmd > max_val_ptpvcmd_vs_pe:  # max value of command PTPVCMD for all iter.
                max_val_ptpvcmd_vs_pe = calc_max_val_ptpvcmd

            if calc_max_val_icmd > max_val_icmd:  # calc max value of ICMD for all iter.
                max_val_icmd = calc_max_val_icmd

            if (calc_min_val_icmd < min_val_icmd) and (calc_min_val_icmd > 0):  # calc positive and min value of ICMD
                min_val_icmd = calc_min_val_icmd

        self.max_velocity = int(max_val_ptpvcmd_vs_pe)  # for plot settling
        if int(max_val_ptpvcmd_vs_pe) == 0:  # check if command is zero (for prevent divide by zero)
            self.factor_ptpvcmd_pe = 1
            self.factor_ptpvcmd_icmd = 0.01

        # average(max_pe, min_pe)/max_PTPVCMD command
        self.factor_ptpvcmd_pe = round(((max_val_pe + min_val_pe) / 2) / (int(max_val_ptpvcmd_vs_pe)), 4)
        self.factor_ptpvcmd_icmd = round(((max_val_icmd + min_val_icmd) / 2) / (int(max_val_ptpvcmd_vs_pe)), 4)
        '''
        print("\nmax value PTPTVCMD: ", max_val_ptpvcmd_vs_pe)
        print("max value ICMD: ", max_val_icmd)
        print("min value ICMD: ", min_val_icmd)
        print("factor ptpvcmd for ICMD: ", self.factor_ptpvcmd_icmd)

        print("\nmax value PE: ", max_val_pe)
        print("min value PE: ", min_val_pe)
        print("mean value PE: ", (max_val_pe + min_val_pe)/2)
        print("factor ptpvcmd: ", self.factor_ptpvcmd_pe)
        '''


class Record(Main):

    def samplefac(self, pre_points, acc, velocity, distance, points):

        arr = [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 18, 32, 64]
        pre_points = int(pre_points)
        acc = float(acc)
        distance = float(distance)
        velocity = float(velocity)
        points = int(points)
        self.sample_value = arr[0]
        setling_time = 300

        sum1 = (2 * velocity * 1000) / acc
        sum2 = (distance * 60 * 1000) / velocity

        # print("sum1: " + str(sum1))
        # print("sum2: " + str(sum2))

        profile_time = ((sum1 + sum2) + (pre_points*0.03125*self.sample_value)) + 300
        total_time = 2000*0.03125*self.sample_value
        # print(total_time)
        # print(profile_time)

        i = 0
        while len(arr):
            if total_time < profile_time:
                i += 1
                # print("total time: " + str(total_time))
                # print("profile time: " + str(profile_time))
                self.sample_value = arr[i]
                profile_time = ((sum1 + sum2) + (pre_points * 0.03125 * self.sample_value)) + 300
                total_time = 2000 * 0.03125 * self.sample_value
            else:
                # print("total time: " + str(total_time))
                # print("profile time: " + str(profile_time))
                self.sample_value = arr[i]
                break

        # print(self.sample_value)
        return str(self.sample_value)

    def write_all_record_data(self, get, iteration, write_destination):
        inx_ptpvcmd = 0  # PTPVCMD initialization
        inx_icmd = 1  # ICMD initialization
        inx_pe = 2

        arr_ptpvcmd = array('f', [])
        arr_pe = array('f', [])
        arr_icmd = array('f', [])

        arr = get[get.index("PE") + len("PE") + 1:]
        string = arr.replace(',', ' ')
        columns = string.split()

        for x in range(int(len(columns) / self.num_of_rec_objects)):

            arr_ptpvcmd.insert(x, float(columns[inx_ptpvcmd]))
            arr_pe.insert(x, float(columns[inx_pe]))

            arr_icmd.insert(x, float(columns[inx_icmd]))

            inx_ptpvcmd += self.num_of_rec_objects
            inx_pe += self.num_of_rec_objects
            inx_icmd += self.num_of_rec_objects

        if write_destination == "counts":

            self.MatrixPEAll[iteration][:] = arr_ptpvcmd
            self.MatrixPEAll[iteration + 1][:] = arr_pe

            self.MatrixICMDAll[iteration][:] = arr_ptpvcmd
            self.MatrixICMDAll[iteration + 1][:] = arr_icmd

        if write_destination == "degree":

            self.MatrixPEAll_deg[iteration][:] = arr_ptpvcmd
            self.MatrixPEAll_deg[iteration + 1][:] = arr_pe

    def write_dec_record_data(self, get, iteration, write_destination):
        inx_ptpvcmd = 0  # PTPVCMD initialization
        inx_icmd = 1  # ICMD initialization
        inx_pe = 2

        arr_ptpvcmd = array('f', [])
        arr_pe = array('f', [])
        arr_icmd = array('f', [])

        arr = get[get.index("PE") + len("PE") + 1:]
        string = arr.replace(',', ' ')
        columns = string.split()

        for x in range(int(len(columns) / self.num_of_rec_objects)):
            arr_ptpvcmd.insert(x, float(columns[inx_ptpvcmd]))
            arr_pe.insert(x, float(columns[inx_pe]))

            arr_icmd.insert(x, float(columns[inx_icmd]))

            inx_ptpvcmd += self.num_of_rec_objects
            inx_pe += self.num_of_rec_objects
            inx_icmd += self.num_of_rec_objects

        if write_destination == "counts":
            self.MatrixPEDec[iteration][:] = arr_ptpvcmd
            self.MatrixPEDec[iteration + 1][:] = arr_pe

            self.MatrixICMDDec[iteration][:] = arr_ptpvcmd
            self.MatrixICMDDec[iteration + 1][:] = arr_icmd

        if write_destination == "degree":
            self.MatrixPEDec_deg[iteration][:] = arr_ptpvcmd
            self.MatrixPEDec_deg[iteration + 1][:] = arr_pe

    def write_record_to_file(self, get, file_name):
        recall_file = ""
        try:
            recall_file = open(self.dir_iter + '\\' + file_name, "w")
            recall_file.write(get)

        except ValueError:
            try:
                recall_file = open(self.dir_iter + '\\' + file_name, "w")
                recall_file.write(get)

            finally:
                recall_file.close()

    def record_all_profile(self, iteration):

        rec_param = "ptpvcmd" + " \"" + "icmd" + " \"" + "pe"
        rec_trig_param = "ptpvcmd"
        rec_trig_cond = "10 100 1"
        strec3 = " \"" + rec_param
        strec1 = "record " + self.samplefac("100", self.traj_array[0], self.traj_array[1],
                                            self.traj_array[2], str(self.time_rec)) + " " + \
                 str(self.time_rec) + strec3
        strec2 = "rectrig " + "\"" + rec_trig_param + " " + rec_trig_cond

        PortSerial.port("getmode 0", self.nameComport, 1)
        PortSerial.port("unitsrotpos 0", self.nameComport, 1)
        PortSerial.port("acc " + self.traj_array[0], self.nameComport, 1)
        PortSerial.port("dec " + self.traj_array[0], self.nameComport, 1)
        PortSerial.port("recoff", self.nameComport, 1)
        PortSerial.port(strec1, self.nameComport, 1)
        PortSerial.port(strec2, self.nameComport, 1)

        PortSerial.port("en", self.nameComport, 1)
        PortSerial.port("moveinc " + self.traj_array[2] + " " + self.traj_array[1], self.nameComport, 1)

        cntr = 0

        while int(PortSerial.port("recdone", self.nameComport, 0)) != 1:

            time.sleep(0.2)
            cntr += 1
            if int(PortSerial.port("recdone", self.nameComport, 0)) == 1:
                print("Record All Area -> Done...")

            if cntr == 50:
                print("Record All Fail!")
                break

        if self.mech_setup == 1:
            PortSerial.port("moveabs 0 100", self.nameComport, 1)
            while int(PortSerial.port("stopped", self.nameComport, 0)) != 2:
                time.sleep(0.1)

        PortSerial.port("unitsrotpos 1", self.nameComport, 1)
        get = PortSerial.port("get", self.nameComport, 0)
        PortSerial.port("k", self.nameComport, 1)
        PortSerial.port("unitsrotpos 2", self.nameComport, 1)  # Convert to degree
        get_degree = PortSerial.port("get", self.nameComport, 0)

        self.write_all_record_data(get, iteration, "counts")
        self.write_all_record_data(get_degree, iteration, "degree")
        self.write_record_to_file(get, 'RecordAll.txt')
        self.write_record_to_file(get_degree, 'RecordAll_deg.txt')

    def record_dec_profile(self, iteration):

        rec_param = " \"" + "ptpvcmd" + " \"" + "icmd" + " \"" + "pe"
        rec_trig_param = "ptpvcmd"
        strec1 = "record " + self.samplefac("100", self.traj_array[0], self.traj_array[1],
                                            self.traj_array[2], str(self.time_rec)) + " " + str(self.time_rec)\
                 + rec_param

        PortSerial.port("recoff", self.nameComport, 1)
        PortSerial.port(strec1, self.nameComport, 1)
        rec_trig_cond = "200 300 0"
        PortSerial.port("rectrig " + "\"" + rec_trig_param + " " + rec_trig_cond, self.nameComport, 1)
        PortSerial.port("en", self.nameComport, 1)
        PortSerial.port("unitsrotpos 0", self.nameComport, 1)
        PortSerial.port("moveinc " + self.traj_array[2] + " " + self.traj_array[1], self.nameComport, 1)
        cntr = 0

        while True:

            if int(PortSerial.port("recdone", self.nameComport, 0)) == 1:
                print("Record Dec area -> Done...")
                break

            time.sleep(0.2)
            cntr += 1

            if cntr == 50:
                print("Record Dec fail!")
                break

        PortSerial.port("unitsrotpos 1", self.nameComport, 1)
        get = PortSerial.port("get", self.nameComport, 0)
        PortSerial.port("unitsrotpos 2", self.nameComport, 1)  # Convert to degree
        get_degree = PortSerial.port("get", self.nameComport, 0)
        PortSerial.port("k", self.nameComport, 1)

        self.write_dec_record_data(get, iteration, "counts")
        self.write_dec_record_data(get_degree, iteration, "degree")
        self.write_record_to_file(get, 'RecordDec.txt')
        self.write_record_to_file(get_degree, 'RecordDec_deg.txt')

    def rec_external_iteration(self, iteration):

        try:

            rec_param = "ptpvcmd" + " \"" + "icmd" + " \"" + "pe"
            rec_trig_param = "ptpvcmd"
            rec_trig_cond = "10 100 1"
            strec3 = " \"" + rec_param
            strec1 = "record " + self.ext_traj_calc() + " " + str(self.time_rec) + strec3
            # print(strec1)
            strec2 = "rectrig " + "\"" + rec_trig_param + " " + rec_trig_cond
            # print(strec2)
            PortSerial.port("getmode 0", self.nameComport, 1)
            PortSerial.port("unitsrotpos 0", self.nameComport, 1)
            PortSerial.port("recoff", self.nameComport, 1)
            PortSerial.port(strec1, self.nameComport, 1)
            PortSerial.port(strec2, self.nameComport, 1)

            cntr = 0
            while int(PortSerial.port("recdone", self.nameComport, 0)) != 1:

                time.sleep(0.2)
                cntr += 1
                if int(PortSerial.port("recdone", self.nameComport, 0)) == 1:
                    print("Record All Area -> Done...")

                if cntr == 50:
                    print("Record All Fail!")
                    break

            PortSerial.port("unitsrotpos 1", self.nameComport, 1)
            get = PortSerial.port("get", self.nameComport, 0)
            PortSerial.port("unitsrotpos 2", self.nameComport, 1)
            get_degree = PortSerial.port("get", self.nameComport, 0)
            # print(get)

            self.write_all_record_data(get, iteration, "counts")
            self.write_all_record_data(get_degree, iteration, "degree")
            self.write_record_to_file(get, 'RecordAll.txt')
            self.write_record_to_file(get_degree, 'RecordAll_deg.txt')

            PortSerial.port("recoff", self.nameComport, 1)
            PortSerial.port(strec1, self.nameComport, 1)
            rec_trig_cond = "200 300 0"
            PortSerial.port("rectrig " + "\"" + rec_trig_param + " " + rec_trig_cond, self.nameComport, 1)
            PortSerial.port("unitsrotpos 0", self.nameComport, 1)
            cntr = 0

            while True:

                if int(PortSerial.port("recdone", self.nameComport, 0)) == 1:
                    print("Record Dec area -> Done...")
                    break

                time.sleep(0.2)
                cntr += 1

                if cntr == 50:
                    print("Record Dec fail!")
                    break

            PortSerial.port("unitsrotpos 1", self.nameComport, 1)
            get = PortSerial.port("get", self.nameComport, 0)
            PortSerial.port("unitsrotpos 2", self.nameComport, 1)
            get_degree = PortSerial.port("get", self.nameComport, 0)

            self.write_dec_record_data(get, iteration, "counts")
            self.write_dec_record_data(get_degree, iteration, "degree")
            self.write_record_to_file(get, 'RecordDec.txt')
            self.write_record_to_file(get_degree, 'RecordDec_deg.txt')

        except ValueError:
            print("Something wrong in data record")
            self.rec_external_iteration(iteration)

    def ext_traj_calc(self):

        arr_ptpvcmd = array('f', [])
        arr_pe = array('f', [])
        arr_icmd = array('f', [])

        inx_pcmd = 0
        inx_ptpvcmd = 1

        PortSerial.port("getmode 0", self.nameComport, 1)
        PortSerial.port("unitsrotpos 0", self.nameComport, 1)
        PortSerial.port("recoff", self.nameComport, 1)
        PortSerial.port("record 32 2000 \"pcmd \"ptpvcmd", self.nameComport, 1)
        PortSerial.port("rectrig \"ptpvcmd 10 100 1", self.nameComport, 1)

        cntr = 0
        while int(PortSerial.port("recdone", self.nameComport, 0)) != 1:
            time.sleep(0.2)
            cntr += 1
            if int(PortSerial.port("recdone", self.nameComport, 0)) == 1:
                print("Record Recognise -> Done...")

            if cntr == 50:
                print("Record Recognise Fail!")
                break

        get = PortSerial.port("get", self.nameComport, 0)
        arr = get[get.index("PTPVCMD") + len("PTPVCMD") + 1:]
        string = arr.replace(',', ' ')
        columns = string.split()

        for x in range(int(len(columns) / 2)):
            arr_ptpvcmd.insert(x, float(columns[inx_ptpvcmd]))
            arr_pe.insert(x, float(columns[inx_pcmd]))
            inx_ptpvcmd += 2
            inx_pcmd += 2

        vel_val = int(max(arr_ptpvcmd))
        dist_val = round(abs(float(max(arr_pe)) - float(min(arr_pe))), 2)
        print("Velocity is: ", vel_val, "[rpm]")
        print("max: ", float(max(arr_pe)), "min: ", float(min(arr_pe)))
        print("Distance is: ", dist_val, "[rev]")

        for x in range(int(len(arr_ptpvcmd)) - 1):
            arr_icmd.insert(x, arr_ptpvcmd[x + 1] - arr_ptpvcmd[x])

        acc_val = 1000 * int(max(arr_icmd))
        print("Acceleration is: ", acc_val, "[rpm/s]")
        # print(samplefac("100", acc_val, vel_val, dist_val, "2000"))
        return self.samplefac("100", acc_val, vel_val, dist_val, "2000")


class Support(Main):

    def calc_factor(self, MatrixPE, MatrixICMD):

        max_val_pe = 1
        max_val_icmd = 1
        min_val_pe = 0
        min_val_icmd = 0
        max_val_ptpvcmd_vs_pe = 1

        for ind in range(0, self.numOfIterations * 2, 2):

            calc_max_val_pe = max(np.multiply(self.factorPE, MatrixPE[ind + 1][:]))  # calculation max value of
            # PE for each iteration
            calc_min_val_pe = min(np.multiply(self.factorPE, MatrixPE[ind + 1][:]))  # calculation min value of
            # PE for each iteration
            calc_max_val_ptpvcmd = max(MatrixPE[ind][:])  # calculation max value of PTPVCMD for each iteration

            calc_max_val_icmd = max(MatrixICMD[ind + 1][:])  # calculation max value of ICMD for each iteration
            calc_min_val_icmd = min(MatrixPE[ind + 1][:])  # calculation min value of ICMD for each iteration

            if calc_max_val_pe > max_val_pe:
                max_val_pe = calc_max_val_pe

            if (calc_min_val_pe < min_val_pe) and (calc_min_val_pe > 0):
                min_val_pe = calc_min_val_pe

            if calc_max_val_ptpvcmd > max_val_ptpvcmd_vs_pe:  # max value of command PTPVCMD for all iter.
                max_val_ptpvcmd_vs_pe = calc_max_val_ptpvcmd

            if calc_max_val_icmd > max_val_icmd:  # calc max value of ICMD for all iter.
                max_val_icmd = calc_max_val_icmd

            if (calc_min_val_icmd < min_val_icmd) and (calc_min_val_icmd > 0):  # calc positive and min value of ICMD
                min_val_icmd = calc_min_val_icmd

        self.max_velocity = int(max_val_ptpvcmd_vs_pe)  # for plot settling
        if int(max_val_ptpvcmd_vs_pe) == 0:  # check if command is zero (for prevent divide by zero)
            self.factor_ptpvcmd_pe = 1
            self.factor_ptpvcmd_icmd = 0.01

        # average(max_pe, min_pe)/max_PTPVCMD command
        self.factor_ptpvcmd_pe = round(((max_val_pe + min_val_pe) / 2) / (int(max_val_ptpvcmd_vs_pe)), 4)
        self.factor_ptpvcmd_icmd = round(((max_val_icmd + min_val_icmd) / 2) / (int(max_val_ptpvcmd_vs_pe)), 4)
        '''
        print("\nmax value PTPTVCMD: ", max_val_ptpvcmd_vs_pe)
        print("max value ICMD: ", max_val_icmd)
        print("min value ICMD: ", min_val_icmd)
        print("factor ptpvcmd for ICMD: ", self.factor_ptpvcmd_icmd)

        print("\nmax value PE: ", max_val_pe)
        print("min value PE: ", min_val_pe)
        print("mean value PE: ", (max_val_pe + min_val_pe)/2)
        print("factor ptpvcmd: ", self.factor_ptpvcmd_pe)
        '''

    def mail_send(self):

        print("Sending mail ...")
        msg = MIMEMultipart()
        msg['From'] = "dimagerc@gmail.com"
        msg['To'] = self.address
        msg['Subject'] = "Auto Tune test"
        msg['Date'] = formatdate(localtime=True)
        body = "Tune test is Done"
        # print(body)
        msg.attach(MIMEText(body, 'plain'))
        # print(msg)

        filename = self.directory + '\\Report.docx'

        attachment = open(filename, 'rb')

        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment; filename="Report.docx"')
        msg.attach(part)

        try:

            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(msg['From'], password=self.mail_pass)
            server.sendmail(msg['From'], msg['To'], msg.as_string())
            server.quit()

        except:
            print("Unable to send mail")

    @staticmethod
    def popup(display, time_val):
        root = Tk()
        t1 = Text(root, height=20, width=50)
        s = Scrollbar(root)
        s.pack(side=RIGHT, fill=Y)
        t1.pack(side=LEFT, fill=Y)
        s.config(command=t1.yview)
        t1.config(yscrollcommand=s.set)
        t1.insert(END, display)
        root.after(time_val, lambda: root.destroy())
        mainloop()

    def dump(self):
        dump = PortSerial.port("dump", self.nameComport, 1)
        recset_file = ""
        try:
            recset_file = open(self.dir_iter + '\Dump.txt', "w")
            recset_file.write(dump)
        finally:
            recset_file.close()

    def info(self, table, status):
        infosetup = PortSerial.port("info", self.nameComport, 1)
        srvsnsinfo = PortSerial.port("srvsnsinfo", self.nameComport, 1)
        motorname = PortSerial.port("motorname", self.nameComport, 1)
        recset_file = ""
        table_file = ""
        status_file = ""
        try:
            recset_file = open(self.directory + '\Info.txt', "w")
            recset_file.write(infosetup)
            recset_file.write('============================\n')
            recset_file.write(srvsnsinfo)
            recset_file.write('============================\n')
            recset_file.write(motorname)
        finally:
            recset_file.close()

        try:
            table_file = open(self.dir_iter + '\Hdtunetable.txt', "w")
            table_file.write(table)
        finally:
            table_file.close()

        try:
            status_file = open(self.dir_iter + '\HdtuneSt.txt', "w")
            status_file.write(status)
        finally:
            status_file.close()

    @staticmethod
    def clear_value(value):
        elem = "["
        try:
            num_value = value[:value.index(elem)]

            return float(num_value)
        except ValueError:
            return None

    def dir_iter_make(self, num_iter):

        dir_iteration = self.directory + '\\Iteration_' + num_iter
        # print(dirIter)
        if not os.path.exists(dir_iteration):
            os.makedirs(dir_iteration)
        return dir_iteration

    def write_params_iteration(self, iteration):
        # ********* Write parameters after iteration into array ************
        # ******************************************************************
        for i in range(len(self.parameters)):

            val = self.clear_value(PortSerial.port(self.parameters[i], self.nameComport, 0))
            if str(val) == 'None':
                try:
                    val = float(PortSerial.port(self.parameters[i], self.nameComport, 0))
                except ValueError:
                    val = 0
                    pass

            self.MatrixParams[iteration - 1][i] = float(val)
        # ******************************************************************
        # ******************************************************************

        for inx_add_param in range(len(self.traj_array)):
            self.Matrix_add_out_params[iteration - 1][inx_add_param] = float(self.traj_array[inx_add_param])

        # print(self.MatrixParams)
        # print(self.Matrix_add_out_params)

    def writeparams_total(self):
        # ** Write parameters to file Params.txt ********
        recset_file = open(self.directory + '\\Params.txt', "w")

        for i in range(1, self.numOfIterations + 1, 1):
            recset_file.write("Iteration_" + str(i) + ": ")

            for inxParams in range(0, len(self.parameters), 1):
                recset_file.write(str(self.parameters[inxParams]) + ": ")
                recset_file.write(str(self.MatrixParams[i - 1][inxParams]) + " ")

            recset_file.write("\n")
        recset_file.close()
        # **************************************

    def report(self):
        print("Create report ...")
        parameters_total = self.parameters + self.Matrix_add_name_params
        with open(self.directory + '\\Info.txt', 'r') as f:
            info = f.read()
        f.close()

        document = Document()
        document.add_heading('Tuning: ' + self.name_tune, 0)

        if self.name_tune == 'Advance':
            document.add_paragraph('Trajectory: ' + 'Vcruise= ' +
                                   self.vel + '[rpm]' + ' Distance= ' +
                                   self.distance + '[rev]' + ' Acc/Dec= ' +
                                   self.acc + '[rpm/s]' + '\n')

        try:
            p = document.add_paragraph('System information: \n')
            p.add_run(info).bold = True
        except ValueError:
            pass

        document.add_picture(self.dir_plot + '\\PtpvcmdPe.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\PtpvcmdPe_deg.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\PtpvcmdPeDec_deg.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\PtpvcmdIcmd.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\PtpvcmdPeDec.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\PtpvcmdIcmdDec.png', width=Inches(5))
        document.add_picture(self.dir_plot + '\\Settling.png', width=Inches(5))

        if (self.feedbacktype_val is 12) or (self.feedbacktype_val is 19):
            document.add_paragraph('No factor for PE: \n')
            document.add_picture(self.dir_plot + '\\PtpvcmdPe_nofac.png', width=Inches(5))
            document.add_picture(self.dir_plot + '\\PtpvcmdPeDec_nofac.png', width=Inches(5))
            document.add_picture(self.dir_plot + '\\Settling_nofac.png', width=Inches(5))

        for i in range(len(parameters_total)):
            document.add_picture(self.directory + '\\Plots' + '\\' + parameters_total[i] + '.png', width=Inches(5))

        document.add_page_break()
        document.save(self.directory + "\\Report.docx")
        del document


class Plot(Main):

    def plot_rec_all_pe_degree(self):
        print("Plotting All PE profile record degree ...")
        senceArFactor = self.factorPE
        self.factorPE = 1

        self.calc_factor(self.MatrixPEAll_deg, self.MatrixICMDAll)

        plt.xlim([0, (self.sample_value*31.25*self.time_rec)/1000])
        plt.grid()
        plt.title(
            'PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
            + ' & PE[deg]')

        srr = []
        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index/(self.time_rec/((self.sample_value*31.25*self.time_rec)/1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEAll_deg[ind][:]))
                plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEAll_deg[ind + 1][:]))

            plt.xlabel("Time[ms]")
            plt.ylabel("[degree]")
            plt.savefig(self.dir_plot + "\\PtpvcmdPe_deg.png")
            plt.close()
            self.factorPE = senceArFactor

        except ValueError:
            self.factorPE = senceArFactor
            pass

    def plot_rec_dec_pe_degree(self):
        print("Plotting Dec PE profile record degree ...")
        senceArFactor = self.factorPE
        self.factorPE = 1

        self.calc_factor(self.MatrixPEDec_deg, self.MatrixICMDAll)

        plt.xlim([0, (self.sample_value*31.25*self.time_rec)/1000])
        plt.grid()
        plt.title(
            'PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
            + ' & PE[deg]')

        srr = []
        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index/(self.time_rec/((self.sample_value*31.25*self.time_rec)/1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEDec_deg[ind][:]))
                plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEDec_deg[ind + 1][:]))

            plt.xlabel("Time[ms]")
            plt.ylabel("[degree]")
            plt.savefig(self.dir_plot + "\\PtpvcmdPeDec_deg.png")
            plt.close()
            self.factorPE = senceArFactor

        except ValueError:
            self.factorPE = senceArFactor
            pass

    def plot_rec_all_pe(self):
        print("Plotting All PE profile record ...")
        self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)

        plt.xlim([0, (self.sample_value*31.25*self.time_rec)/1000])
        plt.grid()
        plt.title(
            'PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
            + ' & PE(' + str(self.factorPE) + ')[count]')

        srr = []
        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index/(self.time_rec/((self.sample_value*31.25*self.time_rec)/1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEAll[ind][:]))
                plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEAll[ind + 1][:]))

            plt.xlabel('Time[ms]')
            plt.ylabel('[Counts]')
            plt.savefig(self.dir_plot + '\\PtpvcmdPe.png')
            plt.close()

        except ValueError:
            pass

        #  For feedback's SenceAR and MT SenceAR plottingg without factor
        if (self.feedbacktype_val is 12) or (self.feedbacktype_val is 19):

            # == this part of code with factor 1 for PE ==
            tmp_factor_pe = self.factorPE
            self.factorPE = 1
            self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)
            plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
            plt.grid()
            plt.title(
                'PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
                + ' & PE(' + str(self.factorPE) + ')[count]')

            srr = []
            try:

                for index in range(0, self.time_rec, 1):
                    srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

                for ind in range(0, self.numOfIterations * 2, 2):
                    plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEAll[ind][:]))
                    plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEAll[ind + 1][:]))

                plt.xlabel('Time[ms]')
                plt.ylabel('[Counts]')
                plt.savefig(self.dir_plot + '\\PtpvcmdPe_nofac.png')
                plt.close()
                self.factorPE = tmp_factor_pe

            except ValueError:
                self.factorPE = tmp_factor_pe
                pass

    def plot_rec_all_icmd(self):

        print("Plotting All ICMD profile record ...")
        srr = []

        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_icmd, self.MatrixICMDAll[ind][:]))
                plt.plot(srr[:], self.MatrixICMDAll[ind + 1][:])

            plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
            # plt.figure(1)
            plt.grid()
            plt.title('PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_icmd) +
                      ')' + ' & ICMD[A]')

            plt.xlabel('Time[ms]')
            plt.ylabel('[A]')
            plt.savefig(self.dir_plot + '\\PtpvcmdIcmd.png')
            plt.close()

        except ValueError:
            pass

    def plot_rec_dec_pe(self):

        print("Plotting Dec PE profile record ...")
        self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)

        plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
        plt.grid()
        plt.title('PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
                  + ' & PE(' + str(self.factorPE) + ')[count]')

        srr = []
        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEDec[ind][:]))
                plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEDec[ind + 1][:]))

            plt.xlabel('Time[ms]')
            plt.ylabel('[Counts]')
            plt.savefig(self.dir_plot + '\\PtpvcmdPeDec.png')
            plt.close()

        except ValueError:
            pass

        if (self.feedbacktype_val is 12) or (self.feedbacktype_val is 19):
            # == this part of code with factor 1 for PE ==
            tmp_factor_pe = self.factorPE
            self.factorPE = 1
            self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)
            srr = []
            try:

                for index in range(0, self.time_rec, 1):
                    srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

                plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])

                plt.grid()
                plt.title(
                    'PTPVCMD(' + self.traj_array[1] + '[rpm]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_pe) + ')'
                    + ' & PE(' + str(self.factorPE) + ')')

                for ind in range(0, self.numOfIterations * 2, 2):
                    plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEDec[ind][:]))
                    plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEDec[ind + 1][:]))

                plt.xlabel('Time[ms]')
                plt.ylabel('[Counts]')
                plt.savefig(self.dir_plot + '\\PtpvcmdPeDec_nofac.png')
                plt.close()
                self.factorPE = tmp_factor_pe
            except ValueError:
                self.factorPE = tmp_factor_pe
                pass

        # plt.show()

    def plot_rec_dec_icmd(self):

        print("Plotting Dec ICMD profile record ...")
        srr = []
        try:

            for indx in range(0, self.time_rec, 1):
                srr.append(indx / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_icmd, self.MatrixICMDDec[ind][:]))
                plt.plot(srr[:], self.MatrixICMDDec[ind + 1][:])

            # plt.xlim([0, 1950])
            plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
            # plt.figure(1)
            plt.grid()
            plt.title('PTPVCMD(' + self.traj_array[1] + '[rpm/s]' + ')' + ' *factor(' + str(self.factor_ptpvcmd_icmd) + ')'
                      + ' & ICMD[A]')
            plt.xlabel('Time[ms]')
            plt.ylabel('[A]')
            plt.savefig(self.dir_plot + '\PtpvcmdIcmdDec.png')
            plt.close()
        except ValueError:
            pass
        # plt.show()

    def plot_rec_dec_pe_settling(self):
        print("Plotting Settling PE profile record ...")
        self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)

        plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
        plt.ylim([-float(self.max_velocity * self.factor_ptpvcmd_pe * 0.5),
                  float(self.max_velocity * self.factor_ptpvcmd_pe * 0.5)])
        plt.grid()
        plt.title('Settling plot: PTPVCMD(' + self.traj_array[1] + '[rpm/s]' + ')' + ' *factor(' +
                  str(self.factor_ptpvcmd_pe) + ')' + ' & PE(' + str(self.factorPE) + ')')

        srr = []
        try:

            for index in range(0, self.time_rec, 1):
                srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

            for ind in range(0, self.numOfIterations * 2, 2):
                plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEDec[ind][:]), linewidth=5.0)
                plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEDec[ind + 1][:]))

            plt.xlabel('Time[ms]')
            plt.ylabel('[Counts]')
            plt.savefig(self.dir_plot + '\\Settling.png')
            plt.close()
        except ValueError:
            pass

        if (self.feedbacktype_val is 12) or (self.feedbacktype_val is 19):
            # == this part of code with factor 1 for PE ==
            self.factorPE = 1
            if self.smart_factor is True:
                self.calc_factor(self.MatrixPEAll, self.MatrixICMDAll)

            plt.xlim([0, (self.sample_value * 31.25 * self.time_rec) / 1000])
            plt.ylim([-float(self.max_velocity * self.factor_ptpvcmd_pe * 0.5),
                      float(self.max_velocity * self.factor_ptpvcmd_pe * 0.5)])
            plt.grid()
            plt.title('Settling plot: PTPVCMD(' + self.traj_array[1] + '[rpm/s]' + ')' + ' *factor(' +
                      str(self.factor_ptpvcmd_pe) + ')' + ' & PE(' + str(self.factorPE) + ')')

            srr = []
            try:

                for index in range(0, self.time_rec, 1):
                    srr.append(index / (self.time_rec / ((self.sample_value * 31.25 * self.time_rec) / 1000)))

                for ind in range(0, self.numOfIterations * 2, 2):
                    plt.plot(srr[:], np.multiply(self.factor_ptpvcmd_pe, self.MatrixPEDec[ind][:]), linewidth=5.0)
                    plt.plot(srr[:], np.multiply(self.factorPE, self.MatrixPEDec[ind + 1][:]))

                plt.xlabel('Time[ms]')
                plt.ylabel('[Counts]')
                plt.savefig(self.dir_plot + '\\Settling_nofac.png')
                plt.close()
                self.factorPE = 0.002
            except ValueError:
                pass

    def plot_parameters(self):
        print("Ploting all parameters ...")
        array_value = [0 for e in range(self.numOfIterations)]
        max_array_value = array('f', [])
        min_array_value = array('f', [])
        mean_array_value = array('f', [])
        # ********* Plot values from list commands ************************
        # *****************************************************************
        for i in range(0, len(self.parameters), 1):

            for k in range(0, self.numOfIterations, 1):
                array_value[k] = self.MatrixParams[k][i]

            max_val = np.max(array_value)
            min_val = np.min(array_value)
            mean_val = statistics.mean(array_value)

            for inex in range(len(array_value)):
                array_value[inex] = 0

            max_array_value.insert(i, max_val)
            min_array_value.insert(i, min_val)
            mean_array_value.insert(i, mean_val)

        for indx_parameter in range(0, len(self.parameters), 1):

            diction = {0: "Max: " + (str(max_array_value[indx_parameter])[:5]),
                       1: "Min: " + (str(min_array_value[indx_parameter])[:5]),
                       2: "Mean: " + (str(mean_array_value[indx_parameter])[:5])}
            cntr = -1
            for indx_iteration in range(0, self.numOfIterations, 1):

                if self.numOfIterations < 3:
                    plt.plot(indx_iteration, self.MatrixParams[indx_iteration][indx_parameter], 'bx',
                             label=diction[0])
                    plt.plot(indx_iteration, self.MatrixParams[indx_iteration][indx_parameter], 'bx',
                             label=diction[1])
                    plt.plot(indx_iteration, self.MatrixParams[indx_iteration][indx_parameter], 'bx',
                             label=diction[2])
                else:

                    if indx_iteration < 3:
                        cntr += 1
                        plt.plot(indx_iteration, self.MatrixParams[indx_iteration][indx_parameter], 'bx',
                                 label=diction[cntr])

                    plt.plot(indx_iteration, self.MatrixParams[indx_iteration][indx_parameter], 'bx', label="")

            plt.grid()
            plt.xlim([-1, self.numOfIterations])

            plt.title(self.parameters[indx_parameter])
            legend = plt.legend(loc='center left', bbox_to_anchor=(1, 0.5), shadow=True)

            for label in legend.get_texts():
                label.set_fontsize('large')

            for label in legend.get_lines():
                label.set_linewidth(1.5)  # the legend line width

            plt.savefig(self.dir_plot + '\\' + self.parameters[indx_parameter] + '.png',
                        bbox_extra_artists=(legend,), bbox_inches='tight')
            plt.close()
            # plt.show()
        # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        # |~~~~~~~~~~~~~~~~~~~Plot additional values:~~~AccDec~Vel~Distance~Results~~~~~~~~~~~~~~~~~~~
        # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

        for indx_parameter in range(0, len(self.Matrix_add_name_params), 1):

            for indx_iteration in range(self.numOfIterations):
                # print(self.Matrix_add_out_params[indx_iteration][indx_parameter])
                plt.plot(indx_iteration, self.Matrix_add_out_params[indx_iteration][indx_parameter], 'bx')

            plt.grid()
            # plt.xlim([-1, numOfIterations])
            plt.title(self.Matrix_add_name_params[indx_parameter])
            plt.savefig(self.dir_plot + '\\' + self.Matrix_add_name_params[indx_parameter] + '.png')
            # plt.show()
            plt.close()
        # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~


class Tune(Support, Record, Plot):

    def easy_tune(self):

        display_all = ""

        if self.mech_setup is True:
            PortSerial.port("hometype 35", self.nameComport, 1)
            PortSerial.port("homecmd", self.nameComport, 1)

        for iteration in range(1, self.numOfIterations + 1, 1):

            display_iteration = display_all + "\nIteration " + str(iteration) + \
                                " from " + str(self.numOfIterations) + " ... " + "In Process"

            result = 1
            self.dir_iter = self.dir_iter_make(str(iteration))
            PortSerial.port("clearfaults", self.nameComport, 1)
            PortSerial.port("k", self.nameComport, 1)

            PortSerial.port("hdtune 0", self.nameComport, 1)
            PortSerial.port("hdtune 5", self.nameComport, 1)

            time.sleep(0.2)
            PortSerial.port("en", self.nameComport, 1)
            start_tune_test = time.time()
            status = PortSerial.port("hdtunest 1", self.nameComport, 0)
            status_tmp = status
            print(status)
            table_status_tmp = PortSerial.port("hdtunetable", self.nameComport, 0)
            statusbar = 0
            statusbar_tmp = "0"
            print("Iteration: " + str(iteration))

            self.popup(display_iteration, 5000)
            # =======>>>Loop of the Tuning<<<======================
            while int(statusbar) != 100 and result != 0:

                status_a_t = PortSerial.port("hdtunest", self.nameComport, 0)
                statusbar = PortSerial.port("hdtunebar", self.nameComport, 0)
                table_status = PortSerial.port("hdtunetable", self.nameComport, 0)

                for inx_st in range(len(self.err_code)):  # Catch faults of tuning

                    if status_a_t.find(self.err_code[inx_st]) != -1:
                        print("Tuning Fail cause of: ", self.err_code[inx_st])
                        result = 0
                        break

                if statusbar_tmp != statusbar:
                    statusbar_tmp = statusbar
                    print("Progress Bar: " + statusbar_tmp[2:len(statusbar_tmp) - 2] + " %")

                if table_status_tmp != table_status:
                    table_status_tmp = table_status_tmp + table_status

                if status != status_a_t:
                    status = status_a_t
                    status_tmp = status_tmp + status_a_t
            # =====================================================

            finish_tune_test = time.time() - start_tune_test
            PortSerial.port("k", self.nameComport, 1)

            if result == 0:
                print("Iteration_" + str(iteration) + " is Fail")
                self.array_results.insert(iteration - 1, result)
                self.traj_array = ["0", "0", "0", result]
                display_all += "\nIteration " + str(iteration) + " " + "Fail"
                self.popup(display_all, 2000)
                self.write_params_iteration(iteration)

            else:
                PortSerial.port("unitsrotpos 0", self.nameComport, 0)
                traj = PortSerial.port("hdtunetrajinfo", self.nameComport, 0)
                print("Trajectory of tuning: ", traj)
                split_traj = traj.split()

                try:
                    cut_traj_acc = split_traj[split_traj.index('Acc/Dec:') + 1]
                    cut_traj_vcruise = split_traj[split_traj.index('Vcruise:') + 1]
                    cut_traj_distance = split_traj[split_traj.index('Distance:') + 1]
                    # print("Acceleration: ", cut_traj_acc)
                    # print("Vcruise: ", cut_traj_vcruise)
                    # print("Distance: ", cut_traj_distance)
                except ValueError:
                    print("No find trajectory values")
                    result = 0
                    return result

                self.traj_array = [cut_traj_acc, cut_traj_vcruise, cut_traj_distance, result]

                self.record_all_profile((iteration - 1) * 2)
                self.record_dec_profile((iteration - 1) * 2)

                self.array_results.insert(iteration - 1, result)
                print("Iteration_" + str(iteration) + " is Pass")
                display_all += "\nIteration " + str(iteration) + " " + "Done " + \
                               "Time duration: " + str(round(finish_tune_test/60, 2)) + "[s]"
                self.popup(display_all, 2000)

            self.info(table_status_tmp, status_tmp)
            self.dump()
            self.write_params_iteration(iteration)

        count_zero = 0
        for ix in range(len(self.array_results)):
            if self.array_results[ix] == 0:
                count_zero += 1

        if count_zero == len(self.array_results):  # Check if all iterations are fail
            print("Tests are fail: " + str(count_zero))
            print(">>> Test tuning status: Fail <<<")
            return

        self.writeparams_total()
        # if self.smart_factor is True:
        #    self.calc_factor()
        self.plot_rec_all_pe_degree()
        self.plot_rec_dec_pe_degree()
        self.plot_rec_all_pe()
        self.plot_rec_all_icmd()
        self.plot_rec_dec_pe()
        self.plot_rec_dec_icmd()
        self.plot_rec_dec_pe_settling()
        self.plot_parameters()

        print(">>> Test tuning status: PASS <<<")
        self.report()

        if self.mail_cb is True:
            self.mail_send()

    def advance_tune(self):

        display_all = ""
        for iteration in range(1, self.numOfIterations + 1, 1):

            display_iteration = display_all + "\nIteration " + str(iteration) + \
                                " from " + str(self.numOfIterations) + " ... " + "In Process"

            # number of iteration for tuning, start from 1
            result = 1
            self.dir_iter = self.dir_iter_make(str(iteration))
            PortSerial.port("clearfaults", self.nameComport, 1)
            PortSerial.port("k", self.nameComport, 1)

            PortSerial.port("hdtuneavmode " + self.avmode, self.nameComport, 1)
            PortSerial.port("hdtunenlafrc " + self.ffmode, self.nameComport, 1)
            PortSerial.port("hdtuneigrav " + self.igrav, self.nameComport, 1)

            PortSerial.port("unitsrotpos 0", self.nameComport, 1)

            PortSerial.port("hdtunevcruise " + self.vel + " " + self.vel, self.nameComport, 1)
            PortSerial.port("hdtuneacc " + self.acc + " " + self.acc, self.nameComport, 1)
            PortSerial.port("hdtunedist " + self.distance + " -" + self.distance, self.nameComport, 1)
            self.traj_array = [self.acc, self.vel, self.distance, result]

            if self.sedrive is True:

                PortSerial.port("hdtunereference 0", self.nameComport, 1)
                PortSerial.port("hdtunerefen 0", self.nameComport, 1)
                PortSerial.port("hdtune 0", self.nameComport, 1)
                PortSerial.port("hdtunereference 1", self.nameComport, 1)
                PortSerial.port("hdtunerefen 1", self.nameComport, 1)
                PortSerial.port("hdtune 1", self.nameComport, 1)
            else:
                PortSerial.port("hdtune 0", self.nameComport, 1)
                PortSerial.port("hdtune 6", self.nameComport, 1)

            time.sleep(0.2)
            PortSerial.port("en", self.nameComport, 1)
            start_tune_test = time.time()
            status = PortSerial.port("hdtunest 1", self.nameComport, 0)
            status_tmp = status
            print(status)
            table_status_tmp = PortSerial.port("hdtunetable", self.nameComport, 0)
            statusbar = 0
            statusbar_tmp = "0"
            print("Iteration: " + str(iteration))

            self.popup(display_iteration, 5000)
            # =======>>>Loop of the Tuning<<<======================

            while int(statusbar) != 100 and result != 0:

                status_a_t = PortSerial.port("hdtunest", self.nameComport, 0)
                statusbar = PortSerial.port("hdtunebar", self.nameComport, 0)
                table_status = PortSerial.port("hdtunetable", self.nameComport, 0)

                for inx_st in range(len(self.err_code)):  # Catch faults of tuning

                    if status_a_t.find(self.err_code[inx_st]) != -1:
                        print("Tuning Fail cause of: ", self.err_code[inx_st])
                        result = 0
                        break

                if statusbar_tmp != statusbar:
                    statusbar_tmp = statusbar
                    print("Progress Bar: " + statusbar_tmp[2:len(statusbar_tmp) - 2] + " %")

                if table_status_tmp != table_status:
                    table_status_tmp = table_status_tmp + table_status

                if status != status_a_t:
                    status = status_a_t
                    status_tmp = status_tmp + status_a_t

            # =====================================================
            finish_tune_test = time.time() - start_tune_test
            PortSerial.port("k", self.nameComport, 1)

            if result == 0:
                print("Iteration_" + str(iteration) + " is Fail")
                self.array_results.insert(iteration - 1, result)
                self.traj_array = ["0", "0", "0", result]
                display_all += "\nIteration " + str(iteration) + " " + "Fail"
                self.popup(display_all, 2000)
                self.write_params_iteration(iteration)

            else:

                self.record_all_profile((iteration - 1) * 2)
                self.record_dec_profile((iteration - 1) * 2)

                self.array_results.insert(iteration - 1, result)
                print("Iteration_" + str(iteration) + " is Pass")
                display_all += "\nIteration " + str(iteration) + " " + "Done " + \
                               "Time duration: " + str(round(finish_tune_test/60, 2)) + "[s]"
                self.popup(display_all, 2000)

            self.info(table_status_tmp, status_tmp)
            self.dump()
            self.write_params_iteration(iteration)

        count_zero = 0
        for ix in range(len(self.array_results)):
            if self.array_results[ix] == 0:
                count_zero += 1

        if count_zero == len(self.array_results):
            print("Tests are fail: " + str(count_zero))
            print(">>> Test tuning status: Fail <<<")
            return

        self.writeparams_total()
        # if self.smart_factor is True:
        #    self.calc_factor()
        self.plot_rec_all_pe()
        self.plot_rec_all_icmd()
        self.plot_rec_dec_pe()
        self.plot_rec_dec_icmd()
        self.plot_rec_dec_pe_settling()
        self.plot_parameters()

        print(">>> Test tuning status: PASS <<<")
        self.report()

        if self.mail_cb is True:
            self.mail_send()

    def easy_external_tune(self):
        display_all = ""
        for iteration in range(1, self.numOfIterations + 1, 1):

            display_iteration = display_all + "\nIteration " + str(iteration) + \
                                " from " + str(self.numOfIterations) + " ... " + "In Process"

            # number of iteration for tuning, start from 1
            result = 1
            self.dir_iter = self.dir_iter_make(str(iteration))
            PortSerial.port("clearfaults", self.nameComport, 1)

            PortSerial.port("hdtune 0", self.nameComport, 1)
            PortSerial.port("hdtune 7", self.nameComport, 1)
            time.sleep(0.2)

            start_tune_test = time.time()
            status = PortSerial.port("hdtunest 1", self.nameComport, 0)
            status_tmp = status
            print(status)
            table_status_tmp = PortSerial.port("hdtunetable", self.nameComport, 0)
            statusbar = 0
            statusbar_tmp = "0"
            print("Iteration: " + str(iteration))

            self.popup(display_iteration, 5000)
            # =======>>>Loop of the Tuning<<<======================

            while int(statusbar) != 100 and result != 0:

                status_a_t = PortSerial.port("hdtunest", self.nameComport, 0)
                statusbar = PortSerial.port("hdtunebar", self.nameComport, 0)
                table_status = PortSerial.port("hdtunetable", self.nameComport, 0)

                for inx_st in range(len(self.err_code)):  # Catch faults of tuning

                    if status_a_t.find(self.err_code[inx_st]) != -1:
                        print("Tuning Fail cause of: ", self.err_code[inx_st])
                        result = 0
                        break

                if statusbar_tmp != statusbar:
                    statusbar_tmp = statusbar
                    print("Progress Bar: " + statusbar_tmp[2:len(statusbar_tmp) - 2] + " %")

                if table_status_tmp != table_status:
                    table_status_tmp = table_status_tmp + table_status

                if status != status_a_t:
                    status = status_a_t
                    status_tmp = status_tmp + status_a_t

            # =====================================================

            if result == 0:
                print("Iteration_" + str(iteration) + " is Fail")
                self.array_results.insert(iteration - 1, result)
                self.traj_array = ["0", "0", "0", result]
                display_all += "\nIteration " + str(iteration) + " " + "Fail"
                self.popup(display_all, 2000)
                self.write_params_iteration(iteration)

            else:
                finish_tune_test = time.time() - start_tune_test
                self.rec_external_iteration((iteration - 1) * 2)
                self.array_results.insert(iteration - 1, result)
                print("Iteration_" + str(iteration) + " is Pass")
                display_all += "\nIteration " + str(iteration) + " " + "Done " + \
                               "Time duration: " + str(round(finish_tune_test/60, 2)) + "[s]"
                self.popup(display_all, 2000)

            self.info(table_status_tmp, status_tmp)
            self.dump()
            self.write_params_iteration(iteration)

        count_zero = 0
        for ix in range(len(self.array_results)):
            if self.array_results[ix] == 0:
                count_zero += 1

        if count_zero == len(self.array_results):
            print("Tests are fail: " + str(count_zero))
            print(">>> Test tuning status: Fail <<<")
            return

        self.writeparams_total()
        # if self.smart_factor is True:
        #    self.calc_factor()
        self.plot_rec_all_pe()
        self.plot_rec_all_icmd()
        self.plot_rec_dec_pe()
        self.plot_rec_dec_icmd()
        self.plot_rec_dec_pe_settling()
        self.plot_parameters()

        print(">>> Test tuning status: PASS <<<")
        self.report()

        if self.mail_cb is True:
            self.mail_send()

    def advance_external_tune(self):
        display_all = ""
        for iteration in range(1, self.numOfIterations + 1, 1):

            display_iteration = display_all + "\nIteration " + str(iteration) + \
                                " from " + str(self.numOfIterations) + " ... " + "In Process"
            # number of iteration for tuning, start from 1
            result = 1
            self.dir_iter = self.dir_iter_make(str(iteration))
            PortSerial.port("clearfaults", self.nameComport, 1)

            PortSerial.port("hdtuneavmode " + self.avmode, self.nameComport, 1)
            PortSerial.port("hdtunenlafrc " + self.ffmode, self.nameComport, 1)
            PortSerial.port("hdtuneigrav " + self.igrav, self.nameComport, 1)

            PortSerial.port("hdtune 0", self.nameComport, 1)
            PortSerial.port("hdtune 8", self.nameComport, 1)

            time.sleep(0.2)

            start_tune_test = time.time()
            status = PortSerial.port("hdtunest 1", self.nameComport, 0)
            status_tmp = status
            print(status)
            table_status_tmp = PortSerial.port("hdtunetable", self.nameComport, 0)
            statusbar = 0
            statusbar_tmp = "0"
            print("Iteration: " + str(iteration))

            self.popup(display_iteration, 5000)
            # =======>>>Loop of the Tuning<<<======================

            while int(statusbar) != 100 and result != 0:

                status_a_t = PortSerial.port("hdtunest", self.nameComport, 0)
                statusbar = PortSerial.port("hdtunebar", self.nameComport, 0)
                table_status = PortSerial.port("hdtunetable", self.nameComport, 0)

                for inx_st in range(len(self.err_code)):  # Catch faults of tuning

                    if status_a_t.find(self.err_code[inx_st]) != -1:
                        print("Tuning Fail cause of: ", self.err_code[inx_st])
                        result = 0
                        break

                if statusbar_tmp != statusbar:
                    statusbar_tmp = statusbar
                    print("Progress Bar: " + statusbar_tmp[2:len(statusbar_tmp) - 2] + " %")

                if table_status_tmp != table_status:
                    table_status_tmp = table_status_tmp + table_status

                if status != status_a_t:
                    status = status_a_t
                    status_tmp = status_tmp + status_a_t

            # =====================================================

            if result == 0:
                print("Iteration_" + str(iteration) + " is Fail")
                self.array_results.insert(iteration - 1, result)
                self.traj_array = ["0", "0", "0", result]
                display_all += "\nIteration " + str(iteration) + " " + "Fail"
                self.popup(display_all, 2000)
                self.write_params_iteration(iteration)

            else:
                finish_tune_test = time.time() - start_tune_test
                self.rec_external_iteration((iteration - 1) * 2)
                self.array_results.insert(iteration - 1, result)
                print("Iteration_" + str(iteration) + " is Pass")
                display_all += "\nIteration " + str(iteration) + " " + "Done " + \
                               "Time duration: " + str(round(finish_tune_test/60, 2)) + "[s]"
                self.popup(display_all, 2000)

            self.info(table_status_tmp, status_tmp)
            self.dump()
            self.write_params_iteration(iteration)

        count_zero = 0
        for ix in range(len(self.array_results)):
            if self.array_results[ix] == 0:
                count_zero += 1

        if count_zero == len(self.array_results):
            print("Tests are fail: " + str(count_zero))
            print(">>> Test tuning status: Fail <<<")
            return

        self.writeparams_total()
        if self.smart_factor is True:
            self.calc_factor()
        self.plot_rec_all_pe()
        self.plot_rec_all_icmd()
        self.plot_rec_dec_pe()
        self.plot_rec_dec_icmd()
        self.plot_rec_dec_pe_settling()
        self.plot_parameters()

        print(">>> Test tuning status: PASS <<<")
        self.report()

        if self.mail_cb is True:
            self.mail_send()